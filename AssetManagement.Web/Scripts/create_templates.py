from docx import Document

def replace_all(doc, replacements):
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                full = para.text.replace(old, new)
                if full != para.text:
                    for i, run in enumerate(para.runs):
                        run.text = full if i == 0 else ''
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            full = para.text.replace(old, new)
                            if full != para.text:
                                for i, run in enumerate(para.runs):
                                    run.text = full if i == 0 else ''

import sys
base = sys.argv[1]

# ── SELL TEMPLATE ──
doc = Document(base + r'\AssetManagement.Web\wwwroot\templates\sell_original.docx')
sell_rep = {
    'البائع      :     ':     'البائع      :     شركة بيت الخبرة للاستثمار الاقتصادي (تكنو إنفستمنت)',
    'المشتري :    ':           'المشتري :    {{PARTY_NAME}}',
    'الوحدة  المبيعه : ':     'الوحدة المبيعة: {{ASSET_NAME}} - {{ASSET_LOCATION}}',
    'الثمن الإجمالى   :':     'الثمن الإجمالي: {{AMOUNT}} جنيه مصري',
    'إنه في يوم             الموافق  ': 'إنه في يوم {{CONTRACT_DATE}}',
    'يمتلك الطرف الأول ما هو 00000000000000000000000000000000والمحددة بالحدود الآتية : -':
        'يمتلك الطرف الأول الوحدة {{ASSET_NAME}} الكائنة بـ {{ASSET_LOCATION}} - {{ASSET_CITY}}:',
    'الحـــد البـــــحـــرى : 00000000000000':  'المساحة: {{ASSET_AREA}}',
    'الحـــد الـــــشرقى : 0000000000000000':   'صك الملكية: {{DEED_NUMBER}}',
    'الحد الغربى :  :0000000000000000.':        'رقم القطعة: {{PLOT_NUMBER}}',
    'الحد القبلى :  :000000000000000000':       'رقم العقد: {{CONTRACT_NUMBER}}',
    'ورغبة من الطرف الثاني فى شراء 0000000000000000000000   والكائنة  0000000000000000000، وبعد':
        'ورغبةً من الطرف الثاني شراء {{ASSET_NAME}} الكائنة بـ {{ASSET_LOCATION}}، وبعد',
    '0000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000':
        '{{ASSET_NAME}} - {{ASSET_LOCATION}} - {{ASSET_CITY}} - مساحة {{ASSET_AREA}}',
    'تم هذا البيع نظير ثمن إجمالي موضوع هذا العقد بمبلغ 000000000000000000000000جنيهاً مصرياً (000000000000000):-':
        'تم هذا البيع بثمن إجمالي {{AMOUNT}} جنيه مصري ({{AMOUNT_TEXT}}).',
}
replace_all(doc, sell_rep)
for para in doc.paragraphs:
    if '( طرف ثان مشترى )' in para.text:
        full = '{{PARTY_NAME}} - هوية: {{PARTY_ID}} - هاتف: {{PARTY_PHONE}} (طرف ثانٍ مشترٍ)'
        for i, run in enumerate(para.runs):
            run.text = full if i == 0 else ''
doc.save(base + r'\AssetManagement.Web\wwwroot\templates\sell_template.docx')
print('sell_template.docx OK')

# ── RENT TEMPLATE ──
doc = Document(base + r'\AssetManagement.Web\wwwroot\templates\rent_original.docx')
rent_rep = {
    'إسم المستأجر :':           'إسم المستأجر : {{PARTY_NAME}}',
    '  الوحدة رقم : ':          'الوحدة: {{ASSET_NAME}} - {{ASSET_LOCATION}}',
    'إيجار الشهري للوحدة :':   'إيجار شهري: {{AMOUNT}} جنيه',
    'التأمين :':                 'التأمين: {{SECURITY_DEPOSIT}} جنيه',
    'إنه في يوم          الموافق                     تم تحرير هذا العقد  بحضور كلا من : ':
        'إنه في يوم {{CONTRACT_DATE}} تم تحرير هذا العقد بحضور:',
    'يمتلك الطرف الاول  ماهو 000000000000ورغبه من الطرف الثاني في استئجار 0000000000  بنفس العقار لاستخدامها  شقة  فقد تلاقت  اراده الطرفيـن وتم الاتفاق علي الاتي :':
        'يمتلك الطرف الأول {{ASSET_NAME}} الكائنة بـ {{ASSET_LOCATION}} - {{ASSET_CITY}}، ورغبةً من الطرف الثاني في استئجارها لاستخدامها شقة، اتفق الطرفان على:',
    'أجر الطرف الأول المؤجر إلى الطرف الثانى المستأجر 00000000000000000000000000  ، وذلك وفقا':
        'أجّر الطرف الأول إلى الطرف الثاني {{ASSET_NAME}} - {{ASSET_LOCATION}}، وذلك وفقاً',
    'إتفق الطرفان على أن تكون القيمة الإيجارية الشهرية للوحدة  المؤجره موضوع هذا العقد هى مبلغ وقــدره  000000ج  0000000000شهرياً  شاملة  الضريبة العقارية .':
        'القيمة الإيجارية الشهرية {{AMOUNT}} جنيه مصري ({{AMOUNT_TEXT}}) شاملة الضريبة العقارية.',
    'مدة هذا العقد هى 000000000000000000000000000000000000 غير قابله للتجديد':
        'مدة هذا العقد من {{START_DATE}} حتى {{END_DATE}} غير قابلة للتجديد',
    ' حدد مبلغ التأمين بمبلغ 000000000000000000 ) وقدره,':
        'مبلغ التأمين {{SECURITY_DEPOSIT}} جنيه مصري',
    'انه في 0000000000000000  في تمام الساعه     :       بحضور كلا من :':
        'إنه في {{CONTRACT_DATE}} بحضور كلٍّ من:',
    'استلم الطرف الثاني من الطرف الاول   بموجب التوقيع على هذا المحضر من شركه بيت الخبرة للإستثمار الإقتصادى           (تكنو إنفستمنت ) 00000000000000000000000000000000 بدء من':
        'استلم {{PARTY_NAME}} من شركة بيت الخبرة وحدة {{ASSET_NAME}} بدءاً من',
}
replace_all(doc, rent_rep)
for para in doc.paragraphs:
    if para.text.strip() == 'طرف ثاني مستأجر':
        full = '{{PARTY_NAME}} - هوية: {{PARTY_ID}} - هاتف: {{PARTY_PHONE}} (طرف ثانٍ مستأجر)'
        for i, run in enumerate(para.runs):
            run.text = full if i == 0 else ''
doc.save(base + r'\AssetManagement.Web\wwwroot\templates\rent_template.docx')
print('rent_template.docx OK')

# ── RENT COMMERCIAL TEMPLATE ──
doc = Document(base + r'\AssetManagement.Web\wwwroot\templates\rent_commercial_original.docx')
comm_rep = {
    'المستأجر :':               'المستأجر : {{PARTY_NAME}}',
    'الوحدة المستأجرة: المحل رقم  ': 'الوحدة التجارية: {{ASSET_NAME}} - {{ASSET_LOCATION}}',
    'الإيجار الشهري للوحدة : ': 'الإيجار الشهري: {{AMOUNT}} جنيه',
    'التامين :':                 'التأمين: {{SECURITY_DEPOSIT}} جنيه',
    'انه في يوم                       الموافق                             تحرر هذا العقد بيـن كلا من  :':
        'إنه في يوم {{CONTRACT_DATE}} تحرر هذا العقد بين:',
    'أجر الطرف الأول  المؤجر  بصفته إلى الطرف الثانى  المستأجر ما 0000000000000000000000000000 وقد تم':
        'أجّر الطرف الأول إلى الطرف الثاني الوحدة التجارية {{ASSET_NAME}} - {{ASSET_LOCATION}} - {{ASSET_CITY}}، وقد تم',
    'إتفق الطرفان على أن تكون القيمة الإيجارية الشهرية للمحل المؤجر موضوع هذا العقد هى مبلغ وقــدره 000000000000000000000000':
        'القيمة الإيجارية الشهرية {{AMOUNT}} جنيه مصري ({{AMOUNT_TEXT}}) شهرياً.',
    'مدة هذا العقد هى  00000000000000000000 غير قابله للتجديد':
        'مدة هذا العقد من {{START_DATE}} حتى {{END_DATE}} غير قابلة للتجديد',
    ' تم الاتفاق بين الطرفان على التزام الطرف الثانى (المستأجر ) 0000000000000000بسداد  ( فقط 0000000000000 )للطرف الأول كتأميـن للعقد':
        'مبلغ التأمين {{SECURITY_DEPOSIT}} جنيه مصري للطرف الأول كتأمين للعقد',
    'يتم سداد القيمه الايجاريــــــــــه شهريا بدايه من اليوم الخامس عشر إلى اليوم التاسع عشر من كل شـــهرعلى ان يتم سداد اول قيمة إيجارية بداية من 000000000.':
        'يتم السداد من اليوم 15 حتى 19 من كل شهر، وأول سداد من {{START_DATE}}.',
    'انه في يوم 0000000000 الموافق 00000000000000فى تمام الساعه   :   بحضور كلا من :':
        'إنه في يوم {{CONTRACT_DATE}} بحضور كلٍّ من:',
}
replace_all(doc, comm_rep)
for para in doc.paragraphs:
    if 'طرف ثان مستاجر' in para.text:
        full = '{{PARTY_NAME}} - هوية: {{PARTY_ID}} - هاتف: {{PARTY_PHONE}} (طرف ثانٍ مستأجر)'
        for i, run in enumerate(para.runs):
            run.text = full if i == 0 else ''
doc.save(base + r'\AssetManagement.Web\wwwroot\templates\rent_commercial_template.docx')
print('rent_commercial_template.docx OK')
print('ALL DONE')