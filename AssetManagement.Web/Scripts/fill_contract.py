import sys, json
from docx import Document

def fill(template_path, output_path, data):
    doc = Document(template_path)
    def rep_para(para):
        full = para.text
        for k, v in data.items():
            full = full.replace('{{' + k + '}}', str(v) if v else '')
        if full != para.text:
            for i, run in enumerate(para.runs):
                run.text = full if i == 0 else ''
    for para in doc.paragraphs:
        rep_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    rep_para(para)
    doc.save(output_path)

if __name__ == '__main__':
    fill(sys.argv[1], sys.argv[2], json.loads(sys.argv[3]))
    print('OK')