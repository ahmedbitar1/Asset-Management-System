#!/usr/bin/env python3
"""
prepare_templates.py - Adds placeholders to the Word templates
Run this ONCE to prepare the templates.
"""
import sys
import os
from docx import Document

def replace_text_in_doc(doc, old_text, new_text):
    """Replace text in all paragraphs"""
    for para in doc.paragraphs:
        if old_text in para.text:
            full = para.text.replace(old_text, new_text)
            if para.runs:
                para.runs[0].text = full
                for run in para.runs[1:]:
                    run.text = ''
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        full = para.text.replace(old_text, new_text)
                        if para.runs:
                            para.runs[0].text = full
                            for run in para.runs[1:]:
                                run.text = ''

templates_dir = sys.argv[1] if len(sys.argv) > 1 else '.'

# ======== SELL Template ========
sell_path = os.path.join(templates_dir, 'sell.docx')
if os.path.exists(sell_path):
    doc = Document(sell_path)
    replacements = [
        ('البائع      :     ',   'البائع      :     شركة بيت الخبرة للاستثمار الاقتصادي (تكنو إنفستمنت)'),
        ('المشتري :    ',        'المشتري :    {{PARTY_NAME}}'),
        ('الوحدة  المبيعه : ',   'الوحدة المبيعة : {{ASSET_NAME}} - {{ASSET_LOCATION}} - {{ASSET_CITY}}'),
        ('الثمن الإجمالى   :',   'الثمن الإجمالي : {{AMOUNT}} جنيه مصري'),
        ('إنه في يوم             الموافق  ', 'إنه في يوم {{CONTRACT_DATE}}'),
        ('يمتلك الطرف الأول ما هو 00000000000000000000000000000000والمحددة بالحدود الآتية : -',
         'يمتلك الطرف الأول الوحدة {{ASSET_NAME}} الكائنة بـ {{ASSET_LOCATION}} - {{ASSET_CITY}} - مساحتها {{ASSET_AREA}} - صك الملكية رقم {{DEED_NUMBER}}'),
        ('الحـــد البـــــحـــرى : 00000000000000', 'رقم العقد: {{CONTRACT_NUMBER}}'),
        ('الحـــد الـــــشرقى : 0000000000000000', 'رقم القطعة: {{PLOT_NUMBER}}'),
        ('الحد الغربى :  :0000000000000000.', 'هاتف المشتري: {{PARTY_PHONE}}'),
        ('الحد القبلى :  :000000000000000000', 'رقم الهوية: {{PARTY_ID}}'),
        ('ورغبة من الطرف الثاني فى شراء 0000000000000000000000   والكائنة  0000000000000000000،',
         'ورغبةً من الطرف الثاني شراء {{ASSET_NAME}} الكائنة بـ {{ASSET_LOCATION}} - {{ASSET_CITY}}،'),
        ('0000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000000',
         '{{ASSET_NAME}} - {{ASSET_LOCATION}} - {{ASSET_CITY}} - مساحة {{ASSET_AREA}}'),
        ('تم هذا البيع نظير ثمن إجمالي موضوع هذا العقد بمبلغ 000000000000000000000000جنيهاً مصرياً (000000000000000):-',
         'تم هذا البيع بثمن إجمالي {{AMOUNT}} جنيه مصري ({{AMOUNT_TEXT}}).'),
        ('( طرف ثان مشترى )', '{{PARTY_NAME}} - هوية: {{PARTY_ID}} - هاتف: {{PARTY_PHONE}} (طرف ثانٍ مشترٍ)'),
    ]
    for old, new in replacements:
        replace_text_in_doc(doc, old, new)
    doc.save(os.path.join(templates_dir, 'sell_tpl.docx'))
    print('sell_tpl.docx created')

# ======== RENT Template ========
rent_path = os.path.join(templates_dir, 'rent.docx')
if os.path.exists(rent_path):
    doc = Document(rent_path)
    replacements = [
        ('إسم المستأجر :', 'إسم المستأجر : {{PARTY_NAME}}'),
        ('  الوحدة رقم : ', 'الوحدة: {{ASSET_NAME}} - {{ASSET_LOCATION}} - {{ASSET_CITY}}'),
        ('إيجار الشهري للوحدة :', 'إيجار شهري: {{AMOUNT}} جنيه مصري'),
        ('التأمين :', 'التأمين : {{SECURITY_DEPOSIT}} جنيه'),
        ('إنه في يوم          الموافق                     تم تحرير هذا العقد  بحضور كلا من : ',
         'إنه في يوم {{CONTRACT_DATE}} تم تحرير هذا العقد بحضور:'),
        ('يمتلك الطرف الاول  ماهو 000000000000ورغبه من الطرف الثاني في استئجار 0000000000  بنفس العقار لاستخدامها  شقة  فقد تلاقت  اراده الطرفيـن وتم الاتفاق علي الاتي :',
         'يمتلك الطرف الأول {{ASSET_NAME}} بـ {{ASSET_LOCATION}} - {{ASSET_CITY}}، ورغبةً من الطرف الثاني في استئجارها للسكن، فقد اتفق الطرفان على الآتي:'),
        ('أجر الطرف الأول المؤجر إلى الطرف الثانى المستأجر 00000000000000000000000000  ، وذلك وفقا',
         'أجّر الطرف الأول للطرف الثاني {{ASSET_NAME}} - {{ASSET_LOCATION}}، وذلك وفقاً'),
        ('إتفق الطرفان على أن تكون القيمة الإيجارية الشهرية للوحدة  المؤجره موضوع هذا العقد هى مبلغ وقــدره  000000ج  0000000000شهرياً  شاملة  الضريبة العقارية .',
         'القيمة الإيجارية الشهرية {{AMOUNT}} جنيه مصري ({{AMOUNT_TEXT}}) شهرياً شاملة الضريبة العقارية.'),
        ('مدة هذا العقد هى 000000000000000000000000000000000000 غير قابله للتجديد',
         'مدة هذا العقد من {{START_DATE}} حتى {{END_DATE}} غير قابلة للتجديد'),
        (' حدد مبلغ التأمين بمبلغ 000000000000000000 ) وقدره,',
         'مبلغ التأمين {{SECURITY_DEPOSIT}} جنيه مصري'),
        ('انه في 0000000000000000  في تمام الساعه     :       بحضور كلا من :',
         'إنه في {{CONTRACT_DATE}} بحضور كلٍّ من:'),
        ('طرف ثاني مستأجر',
         '{{PARTY_NAME}} - هوية: {{PARTY_ID}} - هاتف: {{PARTY_PHONE}} (طرف ثانٍ مستأجر)'),
    ]
    for old, new in replacements:
        replace_text_in_doc(doc, old, new)
    doc.save(os.path.join(templates_dir, 'rent_tpl.docx'))
    print('rent_tpl.docx created')

# ======== RENT COMMERCIAL Template ========
comm_path = os.path.join(templates_dir, 'rent_commercial.docx')
if os.path.exists(comm_path):
    doc = Document(comm_path)
    replacements = [
        ('المستأجر :', 'المستأجر : {{PARTY_NAME}}'),
        ('الوحدة المستأجرة: المحل رقم  ', 'الوحدة التجارية: {{ASSET_NAME}} - {{ASSET_LOCATION}} - {{ASSET_CITY}}'),
        ('الإيجار الشهري للوحدة : ', 'الإيجار الشهري: {{AMOUNT}} جنيه مصري'),
        ('التامين :', 'التأمين: {{SECURITY_DEPOSIT}} جنيه'),
        ('انه في يوم                       الموافق                             تحرر هذا العقد بيـن كلا من  :',
         'إنه في يوم {{CONTRACT_DATE}} تحرر هذا العقد بين:'),
        ('أجر الطرف الأول  المؤجر  بصفته إلى الطرف الثانى  المستأجر ما 0000000000000000000000000000 وقد تم',
         'أجّر الطرف الأول للطرف الثاني الوحدة التجارية {{ASSET_NAME}} - {{ASSET_LOCATION}}، وقد تم'),
        ('إتفق الطرفان على أن تكون القيمة الإيجارية الشهرية للمحل المؤجر موضوع هذا العقد هى مبلغ وقــدره 000000000000000000000000',
         'القيمة الإيجارية الشهرية {{AMOUNT}} جنيه مصري ({{AMOUNT_TEXT}}) شهرياً.'),
        ('مدة هذا العقد هى  00000000000000000000 غير قابله للتجديد',
         'مدة هذا العقد من {{START_DATE}} حتى {{END_DATE}} غير قابلة للتجديد'),
        (' تم الاتفاق بين الطرفان على التزام الطرف الثانى (المستأجر ) 0000000000000000بسداد  ( فقط 0000000000000 )للطرف الأول كتأميـن للعقد',
         'مبلغ التأمين {{SECURITY_DEPOSIT}} جنيه للطرف الأول كتأمين للعقد'),
        ('انه في يوم 0000000000 الموافق 00000000000000فى تمام الساعه   :   بحضور كلا من :',
         'إنه في يوم {{CONTRACT_DATE}} بحضور كلٍّ من:'),
        ('طرف ثان مستاجر',
         '{{PARTY_NAME}} - هوية: {{PARTY_ID}} - هاتف: {{PARTY_PHONE}} (طرف ثانٍ مستأجر)'),
    ]
    for old, new in replacements:
        replace_text_in_doc(doc, old, new)
    doc.save(os.path.join(templates_dir, 'rent_commercial_tpl.docx'))
    print('rent_commercial_tpl.docx created')

print('All templates prepared!')